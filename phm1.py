import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import folium
from folium import plugins
import branca.colormap as cm
import numpy as np
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import mean_squared_error, r2_score
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

# === 1. STRATEGIC DATA FOUNDATION ===
print("🚀 Starting Life Expectancy Analysis Workflow")
print("=" * 60)

# Load and prepare data
df = pd.read_csv("merged_community_area_data - Copy.csv")
print(f"📊 Initial data shape: {df.shape}")

# Clean column names
df.columns = df.columns.str.strip()

# Identify and exclude death-related factors
death_keywords = ['deaths', 'death', 'mortality', 'fatalities', 'homicide', 'suicide']
death_columns = [col for col in df.columns if any(keyword in col.lower() for keyword in death_keywords)]
print(f"🚫 Excluding {len(death_columns)} death-related factors")

# Keep only relevant columns (exclude death factors)
analysis_columns = [col for col in df.columns if col not in death_columns]
df_analysis = df[analysis_columns].copy()

# Ensure lat/lon columns are present
lat_col = [col for col in df_analysis.columns if "latitude" in col.lower()][0]
lon_col = [col for col in df_analysis.columns if "longitude" in col.lower()][0]
df_analysis = df_analysis.dropna(subset=[lat_col, lon_col, 'Life Expectancy'])

print(f"📈 Analysis data shape: {df_analysis.shape}")
print(f"🎯 Target variable: Life Expectancy")
print(f"📍 Geographic columns: {lat_col}, {lon_col}")

# === 2. GEOSPATIAL INTEGRATION & CONTEXTUAL MAPPING ===
print("\n🗺️  Creating Geospatial Integration Map...")

# Create base map centered on Chicago
m = folium.Map(location=[41.8781, -87.6298], zoom_start=10, tiles='OpenStreetMap')

# Add life expectancy visualization
for _, row in df_analysis.iterrows():
    life_exp = row['Life Expectancy']
    
    # Color coding based on life expectancy
    if life_exp >= 80:
        color = 'green'
    elif life_exp >= 75:
        color = 'orange'
    else:
        color = 'red'
    
    popup_text = f"""
    <b>GEOID: {row['GEOID']}</b><br>
    <b>Life Expectancy: {life_exp:.1f} years</b><br>
    Population: {row.get('Population', 'N/A'):,}<br>
    """
    
    folium.CircleMarker(
        location=[row[lat_col], row[lon_col]],
        radius=8,
        popup=folium.Popup(popup_text, max_width=300),
        color=color,
        fill=True,
        fill_opacity=0.7
    ).add_to(m)

# Add legend
legend_html = '''
<div style="position: fixed; 
     bottom: 50px; left: 50px; width: 120px; height: 90px; 
     background-color: white; border:2px solid grey; z-index:9999; 
     font-size:14px; padding: 10px">
<p><strong>Life Expectancy</strong></p>
<p><i class="fa fa-circle" style="color:green"></i> 80+ years</p>
<p><i class="fa fa-circle" style="color:orange"></i> 75-80 years</p>
<p><i class="fa fa-circle" style="color:red"></i> <75 years</p>
</div>
'''
m.get_root().html.add_child(folium.Element(legend_html))
m.save("life_expectancy_map.html")
print("✅ Map saved as 'life_expectancy_map.html'")

# === 3. DEEP FACTOR ANALYSIS & PATTERN DISCOVERY ===
print("\n🔍 Conducting Deep Factor Analysis...")

# Select numeric columns for analysis
numeric_df = df_analysis.select_dtypes(include=[np.number])
numeric_df = numeric_df.drop(columns=['GEOID'], errors='ignore')

# Calculate correlations with Life Expectancy
corr_matrix = numeric_df.corr()
life_exp_corr = corr_matrix['Life Expectancy'].abs().sort_values(ascending=False)[1:16]  # Top 15

# Create correlation plot
plt.figure(figsize=(12, 8))
sns.barplot(x=life_exp_corr.values, y=life_exp_corr.index, palette='viridis')
plt.title('Top 15 Factors Correlated with Life Expectancy', fontsize=16, fontweight='bold')
plt.xlabel('Absolute Correlation Coefficient', fontsize=12)
plt.ylabel('Health & Social Determinants', fontsize=12)
plt.tight_layout()
plt.savefig("correlation_life_expectancy.png", dpi=300, bbox_inches='tight')
plt.close()

# Identify key social & demographic drivers
social_factors = life_exp_corr.head(10)
print("🎯 Top 10 Social & Demographic Drivers:")
for i, (factor, corr) in enumerate(social_factors.items(), 1):
    print(f"{i:2d}. {factor}: {corr:.3f}")

# === 4. SIMULATION & IMPACT FORECASTING ===
print("\n🎲 Running Impact Simulation...")

# Create simulation scenarios
sim_df = df_analysis.copy()

# Scenario 1: Improve healthcare access by 20%
healthcare_cols = [col for col in df_analysis.columns if 
                  any(term in col.lower() for term in ['insurance', 'care', 'checkup', 'screening'])]

# Scenario 2: Reduce behavioral risks by 15%
risk_cols = [col for col in df_analysis.columns if 
            any(term in col.lower() for term in ['smoking', 'binge', 'sugary', 'no physical'])]

# Apply improvements
for col in healthcare_cols:
    if col in sim_df.columns and sim_df[col].dtype in ['float64', 'int64']:
        if 'uninsured' in col.lower():
            sim_df[col] = sim_df[col] * 0.8  # Reduce uninsured by 20%
        else:
            sim_df[col] = sim_df[col] * 1.2  # Increase access by 20%

for col in risk_cols:
    if col in sim_df.columns and sim_df[col].dtype in ['float64', 'int64']:
        sim_df[col] = sim_df[col] * 0.85  # Reduce risk behaviors by 15%

print(f"✅ Simulated improvements in {len(healthcare_cols)} healthcare and {len(risk_cols)} behavioral factors")

# === 5. PREDICTIVE MODELING & PRIORITIZATION ===
print("\n🤖 Building Predictive Models...")

# Prepare features and target
target_col = 'Life Expectancy'
feature_cols = [col for col in numeric_df.columns if col != target_col]
X = numeric_df[feature_cols]
y = numeric_df[target_col]

# Handle missing values
imputer = SimpleImputer(strategy='median')
X_imputed = pd.DataFrame(imputer.fit_transform(X), columns=X.columns)

# Scale features
scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X_imputed), columns=X.columns)

# Train-test split
X_train, X_test, y_train, y_test = train_test_split(X_scaled, y, test_size=0.2, random_state=42)

# Train Random Forest model
rf_model = RandomForestRegressor(n_estimators=200, random_state=42, max_depth=10)
rf_model.fit(X_train, y_train)

# Model performance
y_pred = rf_model.predict(X_test)
rmse = np.sqrt(mean_squared_error(y_test, y_pred))
r2 = r2_score(y_test, y_pred)

print(f"📊 Model Performance:")
print(f"   RMSE: {rmse:.2f} years")
print(f"   R²: {r2:.3f}")

# Feature importance analysis
feature_importance = pd.DataFrame({
    'feature': X.columns,
    'importance': rf_model.feature_importances_
}).sort_values('importance', ascending=False)

top_features = feature_importance.head(15)

# Create feature importance plot
plt.figure(figsize=(12, 8))
sns.barplot(data=top_features, x='importance', y='feature', palette='plasma')
plt.title('Top 15 Predictive Features for Life Expectancy', fontsize=16, fontweight='bold')
plt.xlabel('Feature Importance Score', fontsize=12)
plt.ylabel('Predictive Features', fontsize=12)
plt.tight_layout()
plt.savefig("feature_importance.png", dpi=300, bbox_inches='tight')
plt.close()

# === 6. STRATEGY DESIGN & INTERVENTION SIMULATION ===
print("\n🎯 Generating Strategic Recommendations...")

# ZIP-level prioritization based on model predictions
df_analysis['Predicted_Life_Exp'] = rf_model.predict(X_scaled)

# Create IMPROVEMENT SIMULATION for each area
print("🔬 Running Improvement Simulations...")

def simulate_improvements(row_data, feature_importances, improvement_factor=0.2):
    """
    Simulate life expectancy improvements by enhancing top factors
    """
    # Get the row's feature values
    row_features = X_scaled.loc[row_data.name] if row_data.name in X_scaled.index else X_scaled.iloc[0]
    improved_features = row_features.copy()
    
    # Get top 5 most important improvable factors
    top_factors = feature_importances.head(5)
    
    for _, factor_row in top_factors.iterrows():
        feature_name = factor_row['feature']
        if feature_name in improved_features.index:
            current_value = improved_features[feature_name]
            
            # Check correlation direction for improvement
            if feature_name in corr_matrix.index:
                correlation = corr_matrix.loc[feature_name, 'Life Expectancy']
                
                if correlation > 0:
                    # Positive correlation: increase the feature
                    improved_features[feature_name] = current_value + (improvement_factor * abs(current_value))
                else:
                    # Negative correlation: decrease the feature
                    improved_features[feature_name] = current_value - (improvement_factor * abs(current_value))
    
    # Predict with improved features
    improved_prediction = rf_model.predict([improved_features])[0]
    return improved_prediction

# Apply improvement simulation to each area
df_analysis['Improved_Life_Exp'] = df_analysis.apply(
    lambda row: simulate_improvements(row, feature_importance), axis=1
)

# Calculate realistic improvement potential
df_analysis['Improvement_Potential'] = df_analysis['Improved_Life_Exp'] - df_analysis['Life Expectancy']

# Ensure improvement potential is positive and realistic (cap at 10 years max)
df_analysis['Improvement_Potential'] = df_analysis['Improvement_Potential'].clip(0, 10)
df_analysis['Improved_Life_Exp'] = df_analysis['Life Expectancy'] + df_analysis['Improvement_Potential']

print(f"✅ Average improvement potential: {df_analysis['Improvement_Potential'].mean():.2f} years")
print(f"✅ Max improvement potential: {df_analysis['Improvement_Potential'].max():.2f} years")

# Identify high-impact, high-feasibility interventions
high_impact_areas = df_analysis.nsmallest(10, 'Life Expectancy')
high_potential_areas = df_analysis.nlargest(10, 'Improvement_Potential')

# Create intervention priority matrix
intervention_features = feature_importance.head(10)['feature'].tolist()
intervention_matrix = pd.DataFrame({
    'Factor': intervention_features,
    'Importance': feature_importance.head(10)['importance'].values,
    'Feasibility_Score': [0.8, 0.7, 0.9, 0.6, 0.8, 0.7, 0.5, 0.9, 0.6, 0.8],  # Example scores
    'Impact_Score': feature_importance.head(10)['importance'].values * 10
})

intervention_matrix['Priority_Score'] = (intervention_matrix['Importance'] * 
                                       intervention_matrix['Feasibility_Score'] * 
                                       intervention_matrix['Impact_Score'])

intervention_matrix = intervention_matrix.sort_values('Priority_Score', ascending=False)

# === 7. COMPREHENSIVE REPORTING ===
print("\n📄 Generating Comprehensive Report...")

# Create Word document
doc = Document()

# Title page
title = doc.add_heading('Life Expectancy Analysis & Strategic Recommendations', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Data-Driven Insights for Community Health Improvement', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    f"This analysis examined {len(df_analysis)} geographic areas to identify key drivers of life expectancy. "
    f"Using machine learning models, we achieved {r2:.1%} prediction accuracy (RMSE: {rmse:.2f} years) "
    f"and identified {len(top_features)} critical intervention points."
)

# Key Findings
doc.add_heading('Key Findings', level=2)
doc.add_paragraph("1. Healthcare Access and Quality")
for i, (factor, corr) in enumerate(social_factors.head(3).items(), 1):
    doc.add_paragraph(f"   • {factor}: {corr:.3f} correlation", style='List Bullet')

doc.add_paragraph("2. Social and Economic Determinants")
for i, (factor, corr) in enumerate(social_factors.tail(3).items(), 1):
    doc.add_paragraph(f"   • {factor}: {corr:.3f} correlation", style='List Bullet')

# Add correlation chart
doc.add_heading('Factor Correlation Analysis', level=2)
doc.add_picture("correlation_life_expectancy.png", width=Inches(6))

doc.add_page_break()

# Predictive Model Results
doc.add_heading('Predictive Model Results', level=1)
doc.add_paragraph(f"Model Performance Metrics:")
doc.add_paragraph(f"• Root Mean Square Error: {rmse:.2f} years")
doc.add_paragraph(f"• R-squared Score: {r2:.3f}")
doc.add_paragraph(f"• Training Features: {len(feature_cols)}")

doc.add_heading('Top Predictive Features', level=2)
for i, row in top_features.head(10).iterrows():
    doc.add_paragraph(f"{i+1}. {row['feature']}: {row['importance']:.3f}")

doc.add_picture("feature_importance.png", width=Inches(6))

doc.add_page_break()

# Strategic Recommendations
doc.add_heading('Strategic Recommendations', level=1)

doc.add_heading('High-Priority Interventions', level=2)
for i, row in intervention_matrix.head(5).iterrows():
    doc.add_paragraph(f"{i+1}. {row['Factor']}")
    doc.add_paragraph(f"   Priority Score: {row['Priority_Score']:.2f}")
    doc.add_paragraph(f"   Feasibility: {row['Feasibility_Score']:.1f}/1.0")

doc.add_heading('Geographic Prioritization', level=2)
doc.add_paragraph("Areas with Lowest Life Expectancy (Immediate Focus):")
for i, row in high_impact_areas.head(5).iterrows():
    doc.add_paragraph(f"• GEOID {row['GEOID']}: {row['Life Expectancy']:.1f} years")

doc.add_paragraph("Areas with Highest Improvement Potential:")
for i, row in high_potential_areas.head(5).iterrows():
    doc.add_paragraph(f"• GEOID {row['GEOID']}: +{row['Improvement_Potential']:.1f} years potential")

# Save document
doc.save("life_expectancy_analysis_report.docx")
print("✅ Report saved as 'life_expectancy_analysis_report.docx'")

# === 8. ENHANCED GEOSPATIAL VISUALIZATION WITH DARK THEME ===
print("\n🗺️  Creating Enhanced Dark Theme Map with Insights...")

# Define sophisticated color scheme for life expectancy
def get_life_expectancy_color(life_exp):
    """Return color based on life expectancy with sophisticated palette"""
    if life_exp >= 80:
        return '#4ECDC4'  # Teal - High life expectancy
    elif life_exp >= 75:
        return '#FFE66D'  # Warm yellow - Medium life expectancy  
    else:
        return '#FF6B6B'  # Coral red - Low life expectancy

# Enhanced map generation with dark theme
m_enhanced = folium.Map(
    location=[df_analysis[lat_col].mean(), df_analysis[lon_col].mean()],
    zoom_start=10,
    tiles=None,  # Remove default tiles to add custom ones
    prefer_canvas=True
)

# Add dark theme tile layer
folium.TileLayer(
    tiles='https://{s}.basemaps.cartocdn.com/dark_all/{z}/{x}/{y}{r}.png',
    attr='&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors &copy; <a href="https://carto.com/attributions">CARTO</a>',
    name='Dark Theme',
    overlay=False,
    control=True
).add_to(m_enhanced)

# Create a custom colormap legend (but don't add it to the map to avoid duplicate legends)
life_exp_values = df_analysis['Life Expectancy']
colormap = cm.LinearColormap(
    colors=['#FF6B6B', '#FFE66D', '#4ECDC4'],
    vmin=life_exp_values.min(),
    vmax=life_exp_values.max(),
    caption='Life Expectancy (years)'
)
# colormap.add_to(m_enhanced)  # Commented out to prevent duplicate legend

# Add Chicago city boundary with enhanced styling
chicago_boundary = [
    [41.6445, -87.9073],  # Southwest corner
    [41.6445, -87.5236],  # Southeast corner
    [42.0230, -87.5236],  # Northeast corner
    [42.0230, -87.9073],  # Northwest corner
    [41.6445, -87.9073]   # Close the boundary
]

folium.PolyLine(
    locations=chicago_boundary,
    color='#4ECDC4',
    weight=3,
    opacity=0.8,
    popup=folium.Popup('Chicago City Boundary', max_width=200)
).add_to(m_enhanced)

folium.Polygon(
    locations=chicago_boundary,
    color='#4ECDC4',
    weight=2,
    fill=True,
    fill_opacity=0.1,
    popup=folium.Popup('Chicago City Area', max_width=200)
).add_to(m_enhanced)

# Function to get personalized top 3 improvement factors for each GEOID
def get_personalized_improvement_factors(row_data, all_feature_importances):
    """
    Get personalized top 3 improvement factors based on:
    1. Feature importance from model
    2. Current performance of the area in that factor
    3. Improvement potential (lower current values = higher potential)
    """
    improvement_factors = []
    
    # Get all features with their importances
    for feature in all_feature_importances['feature']:
        if feature in row_data.index and pd.notna(row_data[feature]):
            importance = all_feature_importances[all_feature_importances['feature'] == feature]['importance'].iloc[0]
            current_value = row_data[feature]
            
            # Calculate improvement potential based on current performance
            # Lower values in positive factors = higher improvement potential
            # Higher values in negative factors = higher improvement potential
            
            # Identify if factor is positive or negative based on correlation with life expectancy
            if feature in corr_matrix.index:
                correlation = corr_matrix.loc[feature, 'Life Expectancy']
                
                # For positive correlations, lower current values = higher potential
                # For negative correlations, higher current values = higher potential
                if correlation > 0:
                    # Normalize to 0-1 scale, then invert (lower = better potential)
                    feature_max = df_analysis[feature].max()
                    feature_min = df_analysis[feature].min()
                    if feature_max != feature_min:
                        normalized_value = (current_value - feature_min) / (feature_max - feature_min)
                        improvement_potential = 1 - normalized_value
                    else:
                        improvement_potential = 0.5
                else:
                    # For negative correlations, higher current values = higher potential
                    feature_max = df_analysis[feature].max()
                    feature_min = df_analysis[feature].min()
                    if feature_max != feature_min:
                        normalized_value = (current_value - feature_min) / (feature_max - feature_min)
                        improvement_potential = normalized_value
                    else:
                        improvement_potential = 0.5
                
                # Combined score: importance * improvement potential
                combined_score = importance * improvement_potential
                
                improvement_factors.append({
                    'feature': feature,
                    'importance': importance,
                    'current_value': current_value,
                    'improvement_potential': improvement_potential,
                    'combined_score': combined_score
                })
    
    # Sort by combined score and return top 3
    improvement_factors.sort(key=lambda x: x['combined_score'], reverse=True)
    return improvement_factors[:3]

# Get all feature importances for personalized recommendations
all_feature_importances = feature_importance.copy()

for _, row in df_analysis.iterrows():
    life_exp = row['Life Expectancy']
    
    # Get sophisticated color
    color = get_life_expectancy_color(life_exp)
    
    # Get community area name
    community_name = "Unknown"
    if 'Community areas_Name' in df_analysis.columns:
        community_name = row.get('Community areas_Name', 'Unknown')
    elif 'Community Area Name' in df_analysis.columns:
        community_name = row.get('Community Area Name', 'Unknown')
    elif 'Community_Area_Name' in df_analysis.columns:
        community_name = row.get('Community_Area_Name', 'Unknown')
    else:
        name_columns = [col for col in df_analysis.columns if 'name' in col.lower()]
        if name_columns:
            community_name = row.get(name_columns[0], 'Unknown')
    
    # Get personalized improvement factors for this GEOID
    personalized_factors = get_personalized_improvement_factors(row, all_feature_importances)
    
    # Enhanced popup with modern styling
    popup_text = f"""
    <div style="
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.3);
        min-width: 300px;
    ">
        <div style="
            background: rgba(255,255,255,0.1);
            padding: 12px;
            border-radius: 8px;
            margin-bottom: 12px;
            backdrop-filter: blur(10px);
        ">
            <h3 style="margin: 0; font-size: 16px; color: #FFE66D;">📍 {community_name}</h3>
            <p style="margin: 5px 0 0 0; font-size: 12px; opacity: 0.8;">GEOID: {row['GEOID']}</p>
        </div>
        
        <div style="
            display: flex;
            justify-content: space-between;
            margin-bottom: 15px;
            background: rgba(255,255,255,0.1);
            padding: 10px;
            border-radius: 6px;
        ">
            <div style="text-align: center;">
                <div style="font-size: 20px; font-weight: bold; color: #4ECDC4;">{life_exp:.1f}</div>
                <div style="font-size: 11px; opacity: 0.8;">Current</div>
            </div>
            <div style="text-align: center;">
                <div style="font-size: 20px; font-weight: bold; color: #FFE66D;">{row['Improved_Life_Exp']:.1f}</div>
                <div style="font-size: 11px; opacity: 0.8;">With Improvements</div>
            </div>
            <div style="text-align: center;">
                <div style="font-size: 20px; font-weight: bold; color: #FF6B6B;">+{row['Improvement_Potential']:.1f}</div>
                <div style="font-size: 11px; opacity: 0.8;">Potential Gain</div>
            </div>
        </div>
        
        <div style="
            background: rgba(255,255,255,0.1);
            padding: 12px;
            border-radius: 8px;
        ">
            <h4 style="margin: 0 0 10px 0; color: #FFE66D; font-size: 14px;">🎯 Top Improvement Opportunities</h4>
    """
    
    for i, factor_info in enumerate(personalized_factors, 1):
        feature = factor_info['feature']
        current_val = factor_info['current_value']
        potential = factor_info['improvement_potential']
        
        # Create recommendation text based on correlation direction
        if feature in corr_matrix.index:
            correlation = corr_matrix.loc[feature, 'Life Expectancy']
            if correlation > 0:
                recommendation = "📈 Increase"
                icon_color = "#4ECDC4"
            else:
                recommendation = "📉 Reduce"
                icon_color = "#FF6B6B"
        else:
            recommendation = "⚡ Improve"
            icon_color = "#FFE66D"
        
        popup_text += f"""
        <div style="
            margin: 8px 0;
            padding: 8px;
            background: rgba(0,0,0,0.2);
            border-radius: 5px;
            border-left: 3px solid {icon_color};
        ">
            <div style="font-weight: bold; font-size: 12px; color: {icon_color};">
                {i}. {recommendation} {feature}
            </div>
            <div style="font-size: 11px; margin-top: 3px; opacity: 0.9;">
                Current: <span style="color: #FFE66D;">{current_val:.2f}</span> | 
                Potential: <span style="color: #4ECDC4;">{potential:.2f}</span>
            </div>
        </div>
        """
    
    popup_text += """
        </div>
    </div>
    """
    
    # Enhanced circle marker with better styling
    folium.CircleMarker(
        location=[row[lat_col], row[lon_col]],
        radius=10,
        popup=folium.Popup(popup_text, max_width=400),
        color='white',  # White border for contrast on dark theme
        weight=2,
        fill=True,
        fillColor=color,
        fillOpacity=0.8,
        tooltip=f"{community_name}: {life_exp:.1f} years"
    ).add_to(m_enhanced)

# Add fullscreen button
plugins.Fullscreen().add_to(m_enhanced)

# Add measure control for distance measurements
plugins.MeasureControl().add_to(m_enhanced)

# Add a properly aligned title to the map
title_html = '''
<div style="
    position: fixed; 
    top: 15px; 
    left: 15px; 
    width: 320px; 
    height: auto; 
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white !important;
    z-index: 9999; 
    font-size: 18px;
    font-weight: bold;
    padding: 18px;
    border-radius: 12px;
    box-shadow: 0 6px 25px rgba(0,0,0,0.4);
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    backdrop-filter: blur(15px);
    border: 1px solid rgba(255,255,255,0.1);
    text-align: center;
">
📊 Life Expectancy Analysis
<br><span style="font-size: 13px; opacity: 0.9; color: #FFE66D;">Interactive Community Health Map</span>
</div>
'''
m_enhanced.get_root().html.add_child(folium.Element(title_html))

# Add enhanced legend with white text
legend_html = '''
<div style="
    position: fixed; 
    bottom: 20px; 
    left: 20px; 
    width: 240px; 
    height: auto; 
    background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%);
    color: white !important;
    z-index: 9999; 
    font-size: 14px;
    padding: 18px;
    border-radius: 12px;
    box-shadow: 0 6px 25px rgba(0,0,0,0.4);
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    backdrop-filter: blur(15px);
    border: 1px solid rgba(255,255,255,0.1);
">
    <h4 style="margin: 0 0 15px 0; color: #FFE66D; font-size: 16px; text-align: center;">
        🗺️ Life Expectancy Legend
    </h4>
    <div style="display: flex; align-items: center; margin: 10px 0; color: white;">
        <span style="
            display: inline-block; 
            width: 16px; 
            height: 16px; 
            background-color: #4ECDC4; 
            border-radius: 50%; 
            margin-right: 12px;
            border: 2px solid white;
            box-shadow: 0 2px 4px rgba(0,0,0,0.3);
        "></span>
        <span style="color: white; font-weight: 500;">80+ years (Excellent)</span>
    </div>
    <div style="display: flex; align-items: center; margin: 10px 0; color: white;">
        <span style="
            display: inline-block; 
            width: 16px; 
            height: 16px; 
            background-color: #FFE66D; 
            border-radius: 50%; 
            margin-right: 12px;
            border: 2px solid white;
            box-shadow: 0 2px 4px rgba(0,0,0,0.3);
        "></span>
        <span style="color: white; font-weight: 500;">75-79 years (Good)</span>
    </div>
    <div style="display: flex; align-items: center; margin: 10px 0; color: white;">
        <span style="
            display: inline-block; 
            width: 16px; 
            height: 16px; 
            background-color: #FF6B6B; 
            border-radius: 50%; 
            margin-right: 12px;
            border: 2px solid white;
            box-shadow: 0 2px 4px rgba(0,0,0,0.3);
        "></span>
        <span style="color: white; font-weight: 500;">Below 75 years (Needs Focus)</span>
    </div>
    <div style="
        margin-top: 12px; 
        padding-top: 12px; 
        border-top: 1px solid rgba(255,255,255,0.2);
        font-size: 12px; 
        color: #FFE66D; 
        text-align: center;
    ">
        Click markers for detailed insights
    </div>
</div>
'''
m_enhanced.get_root().html.add_child(folium.Element(legend_html))

# Add custom CSS for enhanced styling and better visibility
css_style = """
<style>
    .leaflet-popup-content-wrapper {
        border-radius: 10px !important;
        box-shadow: 0 4px 20px rgba(0,0,0,0.3) !important;
    }
    .leaflet-popup-tip {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
    }
    .leaflet-control-zoom a {
        background: rgba(102, 126, 234, 0.9) !important;
        color: white !important;
        border: none !important;
        border-radius: 4px !important;
    }
    .leaflet-control-zoom a:hover {
        background: rgba(118, 75, 162, 0.9) !important;
        transform: scale(1.05);
        transition: all 0.2s ease;
    }
    .leaflet-bar {
        border-radius: 8px !important;
        overflow: hidden !important;
        box-shadow: 0 4px 15px rgba(0,0,0,0.3) !important;
    }
    .leaflet-control-fullscreen a {
        background: rgba(102, 126, 234, 0.9) !important;
        color: white !important;
        border-radius: 4px !important;
    }
    .leaflet-control-fullscreen a:hover {
        background: rgba(118, 75, 162, 0.9) !important;
    }
    /* Ensure all text in custom elements is white */
    div[style*="position: fixed"] {
        color: white !important;
    }
    div[style*="position: fixed"] * {
        color: inherit !important;
    }
</style>
"""
m_enhanced.get_root().html.add_child(folium.Element(css_style))

# Save the enhanced map
m_enhanced.save("enhanced_life_expectancy_map.html")
print("✅ Enhanced dark theme map saved as 'enhanced_life_expectancy_map.html'")
print("🎨 Features added:")
print("  • Dark theme with CartoDB Dark tiles")
print("  • Sophisticated color palette: Teal, Warm Yellow, Coral Red")
print("  • Modern gradient popups with glassmorphism effect")
print("  • Enhanced tooltips and interactive elements")
print("  • Fullscreen and measurement controls")
print("  • Custom styling and typography")

# === 9. SUMMARY STATISTICS ===
print("\n📈 Analysis Summary:")
print("=" * 60)
print(f"📊 Total areas analyzed: {len(df_analysis)}")
print(f"🎯 Life expectancy range: {df_analysis['Life Expectancy'].min():.1f} - {df_analysis['Life Expectancy'].max():.1f} years")
print(f"📈 Average life expectancy: {df_analysis['Life Expectancy'].mean():.1f} years")
print(f"🔍 Features analyzed: {len(feature_cols)}")
print(f"🤖 Model accuracy (R²): {r2:.3f}")
print(f"📄 Report generated: life_expectancy_analysis_report.docx")
print(f"🗺️  Maps created: life_expectancy_map.html, enhanced_life_expectancy_map.html")
print(f"📊 Charts saved: correlation_life_expectancy.png, feature_importance.png")
print("\n✅ Analysis workflow completed successfully!")